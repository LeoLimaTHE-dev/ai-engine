from pathlib import Path

from docx import Document


def save_docx(
    content: str,
    output_path: str | Path,
    title: str | None = None,
) -> Path:
    path = Path(output_path)

    path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()

    if title:
        document.add_heading(title, level=1)

    document.add_paragraph(content)

    document.save(path)

    return path
