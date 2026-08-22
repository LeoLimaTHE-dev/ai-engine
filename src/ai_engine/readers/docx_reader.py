import mimetypes
from pathlib import Path
from zipfile import ZipFile

from docx import Document

from ai_engine.models import (
    DocumentContent,
    DocumentImage,
    DocumentTable,
)


def read_docx(
    file_path: str | Path,
) -> DocumentContent:
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    if path.suffix.lower() != ".docx":
        raise ValueError(f"Expected a .docx file, got: {path.suffix}")

    document = Document(path)

    paragraphs: list[str] = []

    tables: list[DocumentTable] = []

    images: list[DocumentImage] = []

    # -------------------------
    # Paragraphs
    # -------------------------

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            paragraphs.append(text)

    # -------------------------
    # Tables
    # -------------------------

    for table_index, table in enumerate(
        document.tables,
        start=1,
    ):
        rows: list[list[str]] = []

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]

            rows.append(cells)

        tables.append(
            DocumentTable(
                rows=rows,
                name=f"Table {table_index}",
                source=path.name,
            )
        )

    # -------------------------
    # Embedded images
    # -------------------------

    with ZipFile(path, "r") as docx_zip:
        for item_name in docx_zip.namelist():
            if not item_name.startswith("word/media/"):
                continue

            image_data = docx_zip.read(item_name)

            image_name = Path(item_name).name

            media_type, _ = mimetypes.guess_type(image_name)

            images.append(
                DocumentImage(
                    name=image_name,
                    data=image_data,
                    media_type=media_type,
                )
            )

    # -------------------------
    # Result
    # -------------------------

    return DocumentContent(
        source_path=path,
        text="\n".join(paragraphs),
        tables=tables,
        images=images,
        metadata={
            "format": "docx",
            "filename": path.name,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "image_count": len(images),
        },
    )
