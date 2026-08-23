from pathlib import Path

import pymupdf
import pytest
from docx import Document
from PIL import Image

from ai_engine.models import DocumentContent
from ai_engine.readers import read_docx, read_document, read_documents, read_image, read_pdf
import ai_engine.readers.reader as reader_module


def test_read_docx_extracts_text_table_and_metadata(tmp_path):
    path = tmp_path / "report.docx"
    source = Document()
    source.add_paragraph("First paragraph")
    source.add_paragraph("Second paragraph")
    table = source.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "name"
    table.cell(0, 1).text = "value"
    table.cell(1, 0).text = "alpha"
    table.cell(1, 1).text = "10"
    source.save(path)

    document = read_docx(path)

    assert document.text == "First paragraph\nSecond paragraph"
    assert len(document.tables) == 1
    assert document.tables[0].name == "Table 1"
    assert document.tables[0].source == "report.docx"
    assert document.tables[0].rows == [
        ["name", "value"],
        ["alpha", "10"],
    ]
    assert document.images == []
    assert document.metadata == {
        "format": "docx",
        "filename": "report.docx",
        "paragraph_count": 2,
        "table_count": 1,
        "image_count": 0,
    }


def test_read_pdf_extracts_locally_created_digital_text(tmp_path):
    path = tmp_path / "digital.pdf"
    source = pymupdf.open()
    page = source.new_page()
    expected_text = "This is a digital PDF with enough text for classification."
    page.insert_text((72, 72), expected_text)
    source.set_metadata({"title": "Fixture PDF", "author": "pytest"})
    source.save(path)
    source.close()

    document = read_pdf(path)

    assert document.text == f"[PAGE 1]\n{expected_text}"
    assert document.images == []
    assert document.metadata["format"] == "pdf"
    assert document.metadata["page_count"] == 1
    assert document.metadata["document_type"] == "digital"
    assert document.metadata["is_scanned"] is False
    assert document.metadata["is_mixed"] is False
    assert document.metadata["pages_with_text"] == 1
    assert document.metadata["pages_without_text"] == 0
    assert document.metadata["rendered_page_count"] == 0
    assert document.metadata["title"] == "Fixture PDF"
    assert document.metadata["author"] == "pytest"


def test_read_image_preserves_file_bytes_and_metadata(tmp_path):
    path = tmp_path / "sample.png"
    Image.new("RGB", (3, 2), color=(10, 20, 30)).save(path)
    expected_bytes = path.read_bytes()

    document = read_image(path)

    assert document.source_path == path
    assert document.text == ""
    assert document.tables == []
    assert len(document.images) == 1
    assert document.images[0].name == "sample.png"
    assert document.images[0].data == expected_bytes
    assert document.images[0].media_type == "image/png"
    assert document.metadata == {
        "format": "png",
        "filename": "sample.png",
        "image_count": 1,
    }


def test_read_external_image_preserves_original_filename_exactly(tmp_path):
    path = tmp_path / "Painel CDC.jpeg"
    Image.new("RGB", (3, 2), color=(10, 20, 30)).save(path, format="JPEG")

    document = read_image(path)

    assert document.images[0].name == "Painel CDC.jpeg"
    assert document.filename == "Painel CDC.jpeg"


def test_read_docx_keeps_embedded_image_as_internal_name(tmp_path):
    image_path = tmp_path / "source.png"
    Image.new("RGB", (3, 2), color=(10, 20, 30)).save(image_path)
    docx_path = tmp_path / "report.docx"
    source = Document()
    source.add_picture(str(image_path))
    source.save(docx_path)

    document = read_docx(docx_path)

    assert document.filename == "report.docx"
    assert len(document.images) == 1
    assert document.images[0].name.startswith("image")
    assert document.images[0].name.endswith(".png")
    assert document.images[0].name != image_path.name


def test_read_pdf_keeps_rendered_page_as_internal_name(tmp_path):
    path = tmp_path / "scan.pdf"
    source = pymupdf.open()
    source.new_page()
    source.save(path)
    source.close()

    document = read_pdf(path)

    assert document.filename == "scan.pdf"
    assert [image.name for image in document.images] == ["page_1_render.png"]


@pytest.mark.parametrize(
    ("reader", "filename"),
    [
        (read_docx, "missing.docx"),
        (read_pdf, "missing.pdf"),
        (read_image, "missing.png"),
    ],
)
def test_rich_readers_reject_missing_files(reader, filename, tmp_path):
    with pytest.raises(FileNotFoundError, match="File not found"):
        reader(tmp_path / filename)


@pytest.mark.parametrize(
    ("reader", "filename", "message"),
    [
        (read_docx, "document.txt", "Expected a .docx file"),
        (read_pdf, "document.txt", "Expected a .pdf file"),
        (read_image, "image.txt", "Unsupported image type"),
    ],
)
def test_rich_readers_reject_wrong_extensions(reader, filename, message, tmp_path):
    path = tmp_path / filename
    path.write_bytes(b"content")

    with pytest.raises(ValueError, match=message):
        reader(path)


@pytest.mark.parametrize(
    ("extension", "reader_name"),
    [
        (".txt", "read_text"),
        (".md", "read_markdown"),
        (".markdown", "read_markdown"),
        (".csv", "read_csv"),
        (".docx", "read_docx"),
        (".pdf", "read_pdf"),
        (".xlsx", "read_xlsx"),
        (".xlsm", "read_xlsx"),
        (".png", "read_image"),
        (".jpg", "read_image"),
        (".jpeg", "read_image"),
        (".webp", "read_image"),
        (".bmp", "read_image"),
        (".gif", "read_image"),
        (".tiff", "read_image"),
        (".tif", "read_image"),
    ],
)
def test_read_document_dispatches_each_supported_extension(
    extension,
    reader_name,
    monkeypatch,
    tmp_path,
):
    path = tmp_path / f"input{extension}"
    path.write_bytes(b"fixture")
    calls = []

    def fake_reader(received_path):
        calls.append(received_path)
        return DocumentContent(source_path=received_path, text=reader_name)

    monkeypatch.setattr(reader_module, reader_name, fake_reader)

    result = read_document(path)

    assert calls == [path]
    assert result.source_path == path
    assert result.text == reader_name


def test_read_document_rejects_missing_file(tmp_path):
    with pytest.raises(FileNotFoundError, match="File not found"):
        read_document(tmp_path / "missing.txt")


def test_read_document_rejects_unsupported_extension(tmp_path):
    path = tmp_path / "unsupported.bin"
    path.write_bytes(b"content")

    with pytest.raises(ValueError, match="Unsupported file type: .bin"):
        read_document(path)


def test_read_documents_preserves_input_count_and_order(tmp_path):
    paths = [
        tmp_path / "third.txt",
        tmp_path / "first.txt",
        tmp_path / "second.txt",
    ]

    for index, path in enumerate(paths):
        path.write_text(f"content-{index}", encoding="utf-8")

    documents = read_documents(paths)

    assert len(documents) == len(paths)
    assert [document.source_path for document in documents] == paths
    assert [document.text for document in documents] == [
        "content-0",
        "content-1",
        "content-2",
    ]
