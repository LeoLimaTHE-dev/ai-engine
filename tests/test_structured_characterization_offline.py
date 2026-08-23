import json
import importlib
import os

import pymupdf
import pytest
from docx import Document
from openpyxl import load_workbook

from ai_engine.actions import execute_output, execute_structured_result
from ai_engine.exporters import save_text, save_xlsx_tables
from ai_engine.results import OutputRequest, ResultTable, StructuredResult
from ai_engine.session import ConversationMessage, ConversationSession
from ai_engine.structured import parse_structured_result
from ai_engine.structured_errors import OutputValidationError


chat_module = importlib.import_module("ai_engine.chat")


def parse_object(data):
    return parse_structured_result(json.dumps(data))


def test_plain_text_currently_becomes_text_only_result():
    result = parse_structured_result("  Normal answer  ")

    assert result == StructuredResult(message="Normal answer", outputs=[])


def test_minimal_json_object_currently_parses():
    assert parse_object({"message": "Done"}) == StructuredResult(message="Done")


@pytest.mark.parametrize(
    "raw_response",
    [
        "{invalid json}",
        '{"message": "unfinished"',
        '```json\n{"message": "fenced"}\n```',
        'prefix {"message": "value"}',
        '{"message": "value"} suffix',
    ],
    ids=["invalid", "partial", "fenced", "prefix", "suffix"],
)
def test_non_pure_json_currently_falls_back_to_text(raw_response):
    assert parse_structured_result(raw_response) == StructuredResult(
        message=raw_response,
        outputs=[],
    )


@pytest.mark.parametrize(
    "raw_response",
    ["[1, 2]", '"root string"', "null"],
    ids=["array", "string", "null"],
)
def test_non_object_json_root_currently_falls_back_to_text(raw_response):
    assert parse_structured_result(raw_response) == StructuredResult(
        message=raw_response,
        outputs=[],
    )


def test_empty_json_object_currently_becomes_empty_result():
    assert parse_object({}) == StructuredResult(message="", outputs=[])


def test_missing_message_currently_becomes_empty_string():
    result = parse_object({"outputs": []})

    assert result.message == ""


def test_null_message_currently_becomes_literal_none_string():
    result = parse_object({"message": None})

    assert result.message == "None"


def test_missing_outputs_currently_becomes_empty_list():
    assert parse_object({"message": "Only text"}).outputs == []


@pytest.mark.parametrize("outputs", ["not-a-list", {}, None, 123])
def test_non_list_outputs_are_currently_ignored(outputs):
    assert parse_object({"message": "Done", "outputs": outputs}).outputs == []


def test_non_object_output_items_are_currently_ignored():
    result = parse_object(
        {
            "outputs": [
                "ignored",
                123,
                {"format": "txt", "filename": "kept.txt"},
            ]
        }
    )

    assert result.outputs == [OutputRequest(format="txt", filename="kept.txt")]


def test_missing_format_currently_becomes_empty_string():
    result = parse_object({"outputs": [{"filename": "file"}]})

    assert result.outputs[0].format == ""


def test_null_format_currently_becomes_literal_none_string():
    result = parse_object({"outputs": [{"format": None, "filename": "file"}]})

    assert result.outputs[0].format == "none"


@pytest.mark.parametrize(
    ("output_data", "expected"),
    [
        ({"format": "txt"}, ""),
        ({"format": "txt", "filename": None}, "None"),
    ],
    ids=["missing", "null"],
)
def test_missing_or_null_filename_currently_becomes_string(output_data, expected):
    result = parse_object({"outputs": [output_data]})

    assert result.outputs[0].filename == expected


def test_non_string_content_is_currently_preserved_without_validation():
    result = parse_object(
        {"outputs": [{"format": "txt", "filename": "file.txt", "content": 123}]}
    )

    assert result.outputs[0].content == 123


@pytest.mark.parametrize("tables", ["not-a-list", {}, None, 123])
def test_non_list_tables_are_currently_ignored(tables):
    result = parse_object(
        {"outputs": [{"format": "xlsx", "filename": "file.xlsx", "tables": tables}]}
    )

    assert result.outputs[0].tables == []


def test_non_object_table_items_are_currently_ignored():
    result = parse_object(
        {
            "outputs": [
                {
                    "format": "xlsx",
                    "filename": "file.xlsx",
                    "tables": ["ignored", {"name": "Kept"}],
                }
            ]
        }
    )

    assert result.outputs[0].tables == [ResultTable(name="Kept")]


def test_string_headers_are_currently_split_into_characters():
    result = parse_object(
        {
            "outputs": [
                {
                    "format": "xlsx",
                    "filename": "file.xlsx",
                    "tables": [{"name": "Data", "headers": "AB"}],
                }
            ]
        }
    )

    assert result.outputs[0].tables[0].headers == ["A", "B"]


def test_null_headers_currently_raise_type_error():
    with pytest.raises(TypeError):
        parse_object(
            {
                "outputs": [
                    {
                        "format": "xlsx",
                        "filename": "file.xlsx",
                        "tables": [{"headers": None}],
                    }
                ]
            }
        )


def test_string_rows_are_currently_iterated_and_discarded():
    result = parse_object(
        {
            "outputs": [
                {
                    "format": "xlsx",
                    "filename": "file.xlsx",
                    "tables": [{"rows": "AB"}],
                }
            ]
        }
    )

    assert result.outputs[0].tables[0].rows == []


def test_null_rows_currently_raise_type_error():
    with pytest.raises(TypeError):
        parse_object(
            {
                "outputs": [
                    {
                        "format": "xlsx",
                        "filename": "file.xlsx",
                        "tables": [{"rows": None}],
                    }
                ]
            }
        )


def test_non_list_rows_are_currently_discarded():
    result = parse_object(
        {
            "outputs": [
                {
                    "format": "xlsx",
                    "filename": "file.xlsx",
                    "tables": [{"rows": ["ignored", 123, ["kept"]]}],
                }
            ]
        }
    )

    assert result.outputs[0].tables[0].rows == [["kept"]]


def test_scalar_cells_are_currently_converted_with_str():
    result = parse_object(
        {
            "outputs": [
                {
                    "format": "xlsx",
                    "filename": "file.xlsx",
                    "tables": [{"rows": [[1, True, None, {"key": "value"}]]}],
                }
            ]
        }
    )

    assert result.outputs[0].tables[0].rows == [
        ["1", "True", "None", "{'key': 'value'}"]
    ]


def test_mismatched_header_and_row_widths_are_currently_accepted():
    result = parse_object(
        {
            "outputs": [
                {
                    "format": "xlsx",
                    "filename": "file.xlsx",
                    "tables": [{"headers": ["A", "B"], "rows": [["one"]]}],
                }
            ]
        }
    )

    assert result.outputs[0].tables[0] == ResultTable(
        name="Table",
        headers=["A", "B"],
        rows=[["one"]],
    )


@pytest.mark.parametrize("format_name", ["txt", "TXT", ".txt"])
def test_txt_format_variants_currently_normalize_to_txt(format_name, tmp_path):
    created = execute_output(
        OutputRequest(format=format_name, filename="result", content="content"),
        tmp_path,
    )

    assert created == tmp_path / "result.txt"


@pytest.mark.parametrize("format_name", ["markdown", "csv"])
def test_unsupported_formats_currently_raise_value_error(format_name, tmp_path):
    with pytest.raises(ValueError, match="Unsupported output format"):
        execute_output(
            OutputRequest(format=format_name, filename="result", content="content"),
            tmp_path,
        )


def test_missing_extension_is_currently_appended(tmp_path):
    created = execute_output(
        OutputRequest(format="docx", filename="report", content="content"),
        tmp_path,
    )

    assert created == tmp_path / "report.docx"


def test_conflicting_extension_currently_produces_double_extension(tmp_path):
    created = execute_output(
        OutputRequest(format="pdf", filename="relatorio.docx", content="content"),
        tmp_path,
    )

    assert created == tmp_path / "relatorio.docx.pdf"


@pytest.mark.parametrize(
    ("filename", "expected"),
    [
        ("../../posix.txt", "posix.txt"),
        ("/absolute/path/file.txt", "file.txt"),
    ],
)
def test_posix_paths_are_currently_reduced_to_basename(filename, expected, tmp_path):
    created = execute_output(
        OutputRequest(format="txt", filename=filename, content="content"),
        tmp_path,
    )

    assert created == tmp_path / expected


@pytest.mark.skipif(os.name != "nt", reason="Characterizes WindowsPath semantics")
def test_windows_traversal_is_currently_reduced_to_basename(tmp_path):
    created = execute_output(
        OutputRequest(format="txt", filename=r"..\..\windows.txt", content="content"),
        tmp_path,
    )

    assert created == tmp_path / "windows.txt"


def test_empty_filename_currently_raises_after_extension_only_path_is_built(tmp_path):
    with pytest.raises(ValueError, match="Unsupported output format: $"):
        execute_output(
            OutputRequest(format="txt", filename="", content="content"),
            tmp_path,
        )

    assert not (tmp_path / ".txt").exists()


def test_existing_output_is_currently_overwritten(tmp_path):
    path = tmp_path / "result.txt"
    path.write_text("old", encoding="utf-8")

    created = execute_output(
        OutputRequest(format="txt", filename="result.txt", content="new"),
        tmp_path,
    )

    assert created == path
    assert path.read_text(encoding="utf-8") == "new"


def test_structural_failure_is_now_detected_before_any_output_is_written(
    tmp_path,
):
    result = StructuredResult(
        message="Three outputs",
        outputs=[
            OutputRequest(format="txt", filename="a.txt", content="A"),
            OutputRequest(format="csv", filename="b.csv", content="B"),
            OutputRequest(format="txt", filename="c.txt", content="C"),
        ],
    )

    with pytest.raises(OutputValidationError, match="unsupported format 'csv'"):
        execute_structured_result(result, tmp_path)

    # Preflight planning replaced the historical partial-write behavior.
    assert not (tmp_path / "a.txt").exists()
    assert not (tmp_path / "b.csv").exists()
    assert not (tmp_path / "c.txt").exists()


@pytest.mark.parametrize("extension", ["txt", "md"])
def test_text_exporters_currently_overwrite_existing_files(extension, tmp_path):
    path = tmp_path / f"file.{extension}"
    path.write_text("old", encoding="utf-8")

    save_text("new", path)

    assert path.read_text(encoding="utf-8") == "new"


def test_docx_action_currently_writes_title_and_content_but_ignores_tables(tmp_path):
    created = execute_output(
        OutputRequest(
            format="docx",
            filename="report.docx",
            title="Title",
            content="Body",
            tables=[ResultTable(name="Ignored", headers=["A"], rows=[["cell"]])],
        ),
        tmp_path,
    )

    document = Document(created)
    assert [paragraph.text for paragraph in document.paragraphs] == ["Title", "Body"]
    assert len(document.tables) == 0


def test_pdf_action_currently_writes_text_but_ignores_tables(tmp_path):
    created = execute_output(
        OutputRequest(
            format="pdf",
            filename="report.pdf",
            title="Title",
            content="Body",
            tables=[ResultTable(name="Ignored", headers=["A"], rows=[["cell"]])],
        ),
        tmp_path,
    )

    document = pymupdf.open(created)
    extracted = "\n".join(page.get_text() for page in document)
    document.close()
    assert "Title" in extracted
    assert "Body" in extracted
    assert "Ignored" not in extracted
    assert "cell" not in extracted


def test_xlsx_action_currently_uses_linear_export_without_tables(tmp_path):
    created = execute_output(
        OutputRequest(
            format="xlsx",
            filename="linear.xlsx",
            title="Title",
            content="First\nSecond",
        ),
        tmp_path,
    )

    workbook = load_workbook(created)
    worksheet = workbook["Resultado"]
    assert worksheet["A1"].value == "Title"
    assert worksheet["A3"].value == "First"
    assert worksheet["A4"].value == "Second"
    workbook.close()


def test_xlsx_action_currently_uses_table_export_when_tables_exist(tmp_path):
    created = execute_output(
        OutputRequest(
            format="xlsx",
            filename="tables.xlsx",
            title="Ignored title",
            content="Ignored content",
            tables=[ResultTable(name="Data", headers=["A"], rows=[["cell"]])],
        ),
        tmp_path,
    )

    workbook = load_workbook(created)
    assert workbook.sheetnames == ["Data"]
    assert workbook["Data"]["A1"].value == "A"
    assert workbook["Data"]["A2"].value == "cell"
    workbook.close()


def test_xlsx_tables_currently_accept_mismatched_row_widths(tmp_path):
    path = save_xlsx_tables(
        [ResultTable(name="Data", headers=["A", "B"], rows=[["one"], ["x", "y", "z"]])],
        tmp_path / "mismatched.xlsx",
    )

    workbook = load_workbook(path)
    worksheet = workbook["Data"]
    assert worksheet.max_column == 3
    assert worksheet["B2"].value is None
    assert worksheet["C3"].value == "z"
    workbook.close()


def test_xlsx_sheet_name_is_currently_truncated_to_31_characters(tmp_path):
    long_name = "A" * 40
    path = save_xlsx_tables(
        [ResultTable(name=long_name, rows=[["value"]])],
        tmp_path / "long-name.xlsx",
    )

    workbook = load_workbook(path)
    assert workbook.sheetnames == ["A" * 31]
    workbook.close()


def test_duplicate_xlsx_sheet_names_are_currently_renamed_by_openpyxl(tmp_path):
    path = save_xlsx_tables(
        [ResultTable(name="Same"), ResultTable(name="Same")],
        tmp_path / "duplicate.xlsx",
    )

    workbook = load_workbook(path)
    assert workbook.sheetnames == ["Same", "Same1"]
    workbook.close()


def test_invalid_xlsx_sheet_name_currently_raises_openpyxl_value_error(tmp_path):
    with pytest.raises(ValueError, match="Invalid character"):
        save_xlsx_tables(
            [ResultTable(name="Invalid/Name")],
            tmp_path / "invalid-sheet.xlsx",
        )


def make_chat_session():
    return ConversationSession(provider="openai", documents=[])


def test_chat_with_empty_message_and_outputs_currently_stores_only_user(monkeypatch):
    session = make_chat_session()
    result = StructuredResult(
        message="",
        outputs=[OutputRequest(format="txt", filename="result.txt", content="data")],
    )
    monkeypatch.setattr(chat_module, "run_structured_workflow_documents", lambda **kwargs: result)

    assert chat_module.chat(session, "Create file") is result
    assert session.messages == [ConversationMessage(role="user", content="Create file")]


def test_chat_with_message_and_outputs_currently_stores_both_messages(monkeypatch):
    session = make_chat_session()
    result = StructuredResult(
        message="Report created",
        outputs=[OutputRequest(format="txt", filename="result.txt", content="data")],
    )
    monkeypatch.setattr(chat_module, "run_structured_workflow_documents", lambda **kwargs: result)

    assert chat_module.chat(session, "Create file") is result
    assert session.messages == [
        ConversationMessage(role="user", content="Create file"),
        ConversationMessage(role="assistant", content="Report created"),
    ]


def test_structured_outputs_and_paths_do_not_enter_chat_history_directly(monkeypatch):
    session = make_chat_session()
    result = StructuredResult(
        message="Done",
        outputs=[OutputRequest(format="txt", filename="secret-path.txt", content="data")],
    )
    monkeypatch.setattr(chat_module, "run_structured_workflow_documents", lambda **kwargs: result)

    chat_module.chat(session, "Create file")

    history_text = "\n".join(message.content for message in session.messages)
    assert "secret-path.txt" not in history_text
    assert "data" not in history_text
