import json

import pytest
from docx import Document

from ai_engine.actions import (
    execute_output,
    execute_structured_result,
    sanitize_filename,
)
from ai_engine.results import OutputRequest, ResultTable, StructuredResult
from ai_engine.structured import parse_structured_result


def test_result_models_expose_outputs_and_paths(tmp_path):
    table = ResultTable(
        name="Data",
        headers=["name", "value"],
        rows=[["alpha", "10"]],
    )
    output = OutputRequest(
        format="xlsx",
        filename="report.xlsx",
        title="Report",
        content="Summary",
        tables=[table],
    )
    result = StructuredResult(message="Created", outputs=[output])

    assert table.headers == ["name", "value"]
    assert output.tables == [table]
    assert result.has_outputs is True
    assert result.output_count() == 1
    assert result.output_paths(tmp_path) == [tmp_path / "report.xlsx"]
    assert StructuredResult(message="No files").has_outputs is False


def test_parse_structured_result_parses_multiple_outputs_and_tables():
    raw_response = json.dumps(
        {
            "message": "Files ready",
            "outputs": [
                {
                    "format": "DOCX",
                    "filename": "report.docx",
                    "title": "Report",
                    "content": "Narrative",
                    "tables": [],
                },
                {
                    "format": "XLSX",
                    "filename": "data.xlsx",
                    "tables": [
                        {
                            "name": "Values",
                            "headers": ["item", "amount"],
                            "rows": [["alpha", 10], ["beta", 20]],
                        },
                        "ignored-table",
                    ],
                },
            ],
        }
    )

    result = parse_structured_result(raw_response)

    assert result.message == "Files ready"
    assert result.output_count() == 2
    assert result.outputs[0] == OutputRequest(
        format="docx",
        filename="report.docx",
        title="Report",
        content="Narrative",
    )
    assert result.outputs[1].format == "xlsx"
    assert result.outputs[1].filename == "data.xlsx"
    assert result.outputs[1].tables == [
        ResultTable(
            name="Values",
            headers=["item", "amount"],
            rows=[["alpha", "10"], ["beta", "20"]],
        )
    ]


@pytest.mark.parametrize(
    "raw_response",
    [
        "This is a normal textual response.",
        "[1, 2, 3]",
    ],
)
def test_parse_structured_result_falls_back_to_text(raw_response):
    result = parse_structured_result(f"  {raw_response}  ")

    assert result == StructuredResult(message=raw_response)
    assert result.outputs == []


@pytest.mark.parametrize(
    ("filename", "expected"),
    [
        (r"..\..\resultado.docx", "resultado.docx"),
        (r"folder\report.pdf", "report.pdf"),
        ("simple.txt", "simple.txt"),
    ],
)
def test_sanitize_filename_reduces_paths_to_basename(filename, expected):
    assert sanitize_filename(filename) == expected


def test_execute_output_adds_expected_extension(tmp_path):
    output_dir = tmp_path / "outputs"
    output = OutputRequest(
        format="txt",
        filename="result",
        content="Text content",
    )

    created = execute_output(output, output_dir)

    assert created == output_dir / "result.txt"
    assert created.read_text(encoding="utf-8") == "Text content"


def test_execute_output_keeps_traversal_filename_inside_output_dir(tmp_path):
    output_dir = tmp_path / "safe-output"
    output = OutputRequest(
        format="docx",
        filename=r"..\..\resultado.docx",
        title="Safe title",
        content="Safe content",
    )

    created = execute_output(output, output_dir)

    assert created == output_dir / "resultado.docx"
    assert created.parent.resolve() == output_dir.resolve()
    reopened = Document(created)
    assert [paragraph.text for paragraph in reopened.paragraphs] == [
        "Safe title",
        "Safe content",
    ]


def test_execute_output_rejects_unsupported_format(tmp_path):
    output = OutputRequest(
        format="csv",
        filename="result.csv",
        content="content",
    )

    with pytest.raises(ValueError, match="Unsupported output format: csv"):
        execute_output(output, tmp_path)


def test_execute_structured_result_creates_multiple_outputs(tmp_path):
    output_dir = tmp_path / "multiple"
    result = StructuredResult(
        message="Create both",
        outputs=[
            OutputRequest(format="txt", filename="first.txt", content="First"),
            OutputRequest(format="md", filename="second.md", content="# Second"),
        ],
    )

    created = execute_structured_result(result, output_dir)

    assert created == [
        output_dir / "first.txt",
        output_dir / "second.md",
    ]
    assert created[0].read_text(encoding="utf-8") == "First"
    assert created[1].read_text(encoding="utf-8") == "# Second"

